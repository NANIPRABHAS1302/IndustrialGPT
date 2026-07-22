import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_project_report_docx():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title / Cover Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("IndustrialGPT – Technical Project Report\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("AI for Industrial Knowledge Intelligence: Unified Asset & Operations Brain\nET AI Hackathon Submission\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    doc.add_paragraph("\n" * 3)

    # Metadata Table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    data = [
        ("Project Name", "IndustrialGPT"),
        ("Problem Statement", "PS 8 – AI for Industrial Knowledge Intelligence"),
        ("Participant", "Solo Participant"),
        ("Backend Architecture", "FastAPI, PostgreSQL, ChromaDB, Neo4j, Redis, Celery"),
        ("Frontend Architecture", "React 19, TypeScript, Vite, TailwindCSS")
    ]

    for i, (k, v) in enumerate(data):
        row = table.rows[i]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.text = k
        cell_v.text = v
        set_cell_background(cell_k, "1E293B")
        cell_k.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell_k.paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # Read Project_Report.md and convert headings & text into DOCX
    report_md_path = r"c:\IndustrialGPT\docs\Project_Report.md"
    if os.path.exists(report_md_path):
        with open(report_md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        for line in lines:
            line_str = line.strip()
            if line_str.startswith("# "):
                h = doc.add_heading(line_str[2:], level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif line_str.startswith("## "):
                h = doc.add_heading(line_str[3:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            elif line_str.startswith("### "):
                h = doc.add_heading(line_str[4:], level=3)
                h.runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
            elif line_str.startswith("![") and "](" in line_str:
                # Image embedding
                try:
                    img_rel = line_str.split("](")[1].rstrip(")")
                    img_full = os.path.join(r"c:\IndustrialGPT\docs", img_rel)
                    if os.path.exists(img_full):
                        doc.add_picture(img_full, width=Inches(5.5))
                except Exception as e:
                    print(f"Error adding picture: {e}")
            elif line_str:
                doc.add_paragraph(line_str)

    output_path = r"c:\IndustrialGPT\docs\Project_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated Project_Report.docx -> {output_path}")

if __name__ == "__main__":
    create_project_report_docx()
